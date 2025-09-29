from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from state import ContentOptimizationState
import json
from datetime import datetime

def generate_ing_briefing(state: ContentOptimizationState) -> str:
    """Generate ING Content Briefing in DOCX format"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Content briefing', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Helper function to add table row
    def add_row(table, label, content):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(content) if content else ""
        # Bold the label
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    # Create main table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4.5)
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Veld'
    hdr_cells[1].text = 'Inhoud'
    
    # Add data rows
    add_row(table, 'URL', state.get('url'))
    
    competitor_urls = state.get('competitor_urls', [])
    add_row(table, 'SERP 1#', competitor_urls[0] if len(competitor_urls) > 0 else 'Geen concurrent gevonden')
    add_row(table, 'SERP 2#', competitor_urls[1] if len(competitor_urls) > 1 else 'Geen concurrent gevonden')
    add_row(table, 'SERP 3#', competitor_urls[2] if len(competitor_urls) > 2 else 'Geen concurrent gevonden')
    
    add_row(table, 'Opmerkelijk', state.get('notable_observations'))
    add_row(table, 'Type pagina', state.get('page_type'))
    add_row(table, 'Funnelfase', state.get('funnel_stage'))
    
    # Body copy with instructions
    body_copy_text = f"""{state.get('body_copy_outline', '')}

Schrijf de body copy en verwerk daarin, op een natuurlijke manier, minimaal drie keer het focus keyword en een variatie daarop.

Probeer daarnaast de secundaire keywords en een variatie daarop te verwerken in de tekst. De keywords moeten op een zo natuurlijk mogelijke manier verwerkt worden. 'Keyword stuffing' is niet wenselijk."""
    
    add_row(table, 'Body copy', body_copy_text)
    
    # ING Tone of Voice
    tone_of_voice_text = """Onze communicatie is altijd persoonlijk en begrijpelijk. We klinken actief, brengen lucht en lef in onze teksten en zijn informeel.

Persoonlijk
Zet je klant altijd centraal en verplaats je dus in de klant. Zorg dat je boodschap relevant is. En schrijf en praat altijd inclusief.

Informeel
Blijf sympathiek en innemend. Schrijf en praat eerder informeel dan formeel (je in plaats van u). En gebruik altijd gewone mensentaal, zonder ingewikkeld jargon.

Met lucht en lef
Creëer letterlijk lucht met wit-regels en structuur. En maak keuzes. Neem wat we doen serieus, maar jezelf wat minder. Een grapje mag zeker, als het past.

Begrijpelijk
Bouw je tekst op vanuit 1 hoofdboodschap. Dit helpt je om de tekst beknopt te houden en heldere taal te formuleren. Bij voorkeur zonder jargon. Wees eerlijk en draai er niet omheen.

Actief
Benader je klant altijd positief en denk in oplossingen, niet in problemen. Schrijf energiek en inspireer tot actie."""
    
    add_row(table, 'ING Tone of Voice', tone_of_voice_text)
    
    add_row(table, 'Focus keyword', state.get('focus_keyword'))
    
    secondary_keywords = state.get('secondary_keywords') or []
    add_row(table, 'Secundaire keywords', ', '.join(secondary_keywords) if secondary_keywords else 'Geen secundaire keywords opgegeven')
    
    # Page title suggestion
    page_title_text = f"""Schrijf een page title voor deze pagina die voldoet aan de volgende voorwaarden:
- Max. 60 tekens inclusief spaties
- Gebruik het focus keyword

Voorstel: {state.get('page_title_suggestion', '')}"""
    
    add_row(table, 'Page title suggestie', page_title_text)
    
    # Meta description suggestion
    meta_desc_text = f"""Schrijf een meta description voor deze pagina die voldoet aan de volgende voorwaarden:
- Max. 155 tekens inclusief spaties
- Verwerk het focus keyword
- Probeer een of meerdere secundaire keywords of een variatie daarop te verwerken
- Gebruik een call to action (Bijv. ontdek, bekijk, bestel)

Voorstel: {state.get('meta_description_suggestion', '')}"""
    
    add_row(table, 'Meta description suggestie', meta_desc_text)
    
    # Headers table
    header_suggestions = state.get('header_suggestions', {})
    
    headers_table_text = "Header | Huidig | Nieuw\n"
    headers_table_text += f"H1 | {state.get('title', '')} | {header_suggestions.get('h1', '')}\n"
    
    h2_suggestions = header_suggestions.get('h2', [])
    for i, h2 in enumerate(h2_suggestions[:5], 1):
        headers_table_text += f"H2-{i} | | {h2}\n"
    
    add_row(table, 'Headers inhoudsopgave', headers_table_text)
    
    add_row(table, 'H1 suggestie', header_suggestions.get('h1', ''))
    add_row(table, 'H2 suggestie', '\n'.join(h2_suggestions))
    
    # H3 suggestions
    h3_dict = header_suggestions.get('h3', {})
    h3_text = ""
    for h2, h3_list in h3_dict.items():
        h3_text += f"\n{h2}:\n"
        for h3 in h3_list:
            h3_text += f"  - {h3}\n"
    
    add_row(table, 'H3 suggestie', h3_text)
    
    add_row(table, 'Aanvulling CJE', '')
    add_row(table, 'Inspiratie', f"Zie gedetailleerde analyse in output/report.md")
    
    # Save document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"output/ING_Content_Briefing_{timestamp}.docx"
    doc.save(filename)
    
    return filename